# rag-server/scripts/ingest.py — loads txt, pdf, docx files, splits, embeds, stores in ChromaDB
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from dotenv import load_dotenv
load_dotenv()

from langchain_community.document_loaders import PyPDFLoader, TextLoader, DirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
import docx

DOCS_DIR = os.path.join(os.path.dirname(__file__), "..", "documents")
CHROMA_DIR = os.path.join(os.path.dirname(__file__), "..", "chroma_db")
COLLECTION = "learnai_engineering"

def load_docx(path):
    doc = docx.Document(path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text, metadata={"source": path})]

def load_documents():
    docs = []

    # Load .txt files
    txt_loader = DirectoryLoader(DOCS_DIR, glob="**/*.txt", loader_cls=TextLoader, loader_kwargs={"encoding": "utf-8"}, recursive=True)
    docs.extend(txt_loader.load())

    # Load .pdf files
    pdf_loader = DirectoryLoader(DOCS_DIR, glob="**/*.pdf", loader_cls=PyPDFLoader, recursive=True)
    docs.extend(pdf_loader.load())

    # Load .docx files manually
    for root, _, files in os.walk(DOCS_DIR):
        for f in files:
            if f.endswith(".docx"):
                full_path = os.path.join(root, f)
                try:
                    docs.extend(load_docx(full_path))
                    print(f"  Loaded docx: {f}")
                except Exception as e:
                    print(f"  Failed to load {f}: {e}")

    return docs

def ingest(extra_docs=None):
    print("Loading documents...")
    docs = load_documents()
    if extra_docs:
        docs.extend(extra_docs)
    print(f"Loaded {len(docs)} document(s)")

    if not docs:
        print("No documents found.")
        return 0

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(docs)
    print(f"Split into {len(chunks)} chunks")

    print("Embedding and storing in ChromaDB...")
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    vectorstore = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=COLLECTION,
        persist_directory=CHROMA_DIR,
    )
    print(f"Stored {len(chunks)} chunks in '{COLLECTION}'")
    return len(chunks)

if __name__ == "__main__":
    count = ingest()
    if count:
        print("\nSmoke test — querying: 'What is Newton Second Law?'")
        from langchain_openai import OpenAIEmbeddings
        from langchain_community.vectorstores import Chroma
        embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        vs = Chroma(collection_name=COLLECTION, embedding_function=embeddings, persist_directory=CHROMA_DIR)
        results = vs.similarity_search("What is Newton Second Law?", k=2)
        if results:
            print(f"Retrieved {len(results)} chunk(s):")
            print(results[0].page_content[:300])
        else:
            print("No results.")
