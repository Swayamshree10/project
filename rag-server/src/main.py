# rag-server/src/main.py — FastAPI with /query, /health, and /upload endpoints
import os
import sys
import shutil
import tempfile

sys.path.insert(0, os.path.dirname(__file__))

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from rag_pipeline import get_qa_chain, get_vectorstore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
import docx as python_docx

app = FastAPI(title="Mind Trail RAG Server")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_methods=["*"],
    allow_headers=["*"],
)

chain, retriever = get_qa_chain()

ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}

class Query(BaseModel):
    question: str

@app.get("/health")
def health():
    return {"status": "ok"}

@app.post("/query")
def query(body: Query):
    # Get source docs separately for display
    source_docs = retriever.invoke(body.question)
    answer = chain.invoke(body.question)

    sources = [
        {
            "source": doc.metadata.get("source", "unknown"),
            "snippet": doc.page_content[:200],
        }
        for doc in source_docs
    ]

    print(f"\n[RAG] Question: {body.question}")
    print(f"[RAG] Retrieved {len(sources)} chunks:")
    for i, s in enumerate(sources):
        print(f"  [{i+1}] {s['source']}\n      {s['snippet'][:100]}...")

    return {
        "answer": answer,
        "sources": sources,
        "retrieval_debug": {"chunks_retrieved": len(sources)},
    }

@app.post("/upload")
async def upload(file: UploadFile = File(...), subject: str = "general"):
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{ext}'. Use PDF, TXT, or DOCX.")

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        shutil.copyfileobj(file.file, tmp)
        tmp_path = tmp.name

    try:
        if ext == ".pdf":
            docs = PyPDFLoader(tmp_path).load()
        elif ext == ".txt":
            docs = TextLoader(tmp_path, encoding="utf-8").load()
        elif ext == ".docx":
            doc = python_docx.Document(tmp_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            docs = [Document(page_content=text, metadata={"source": file.filename})]

        for d in docs:
            d.metadata["source"] = file.filename

        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_documents(docs)

        vectorstore = get_vectorstore()
        vectorstore.add_documents(chunks)

        print(f"[Upload] {file.filename} → {len(chunks)} chunks added")
        return {
            "message": f"'{file.filename}' uploaded and indexed successfully.",
            "chunks": len(chunks),
            "subject": subject,
        }
    finally:
        os.unlink(tmp_path)
